from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Preformatted
from reportlab.lib.colors import HexColor
from docx import Document
from docx.shared import Pt, RGBColor
import os

def generate_pdf(report_data: dict, output_path: str, form_input: dict = None):
    try:
        pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
        pdfmetrics.registerFont(TTFont('DejaVu-Bold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
    except:
        pass
    
    doc = SimpleDocTemplate(output_path, pagesize=A4, topMargin=40, bottomMargin=40)
    story = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('CustomTitle', parent=styles['Heading1'], 
                                 fontName='DejaVu-Bold', fontSize=24, textColor=HexColor('#667eea'))
    h1_style = ParagraphStyle('CustomH1', parent=styles['Heading1'], 
                             fontName='DejaVu-Bold', fontSize=16, textColor=HexColor('#333333'))
    h2_style = ParagraphStyle('CustomH2', parent=styles['Heading2'], 
                             fontName='DejaVu-Bold', fontSize=12, textColor=HexColor('#555555'))
    text_style = ParagraphStyle('CustomText', parent=styles['Normal'], 
                                fontName='DejaVu', fontSize=10, leading=14)
    small_style = ParagraphStyle('CustomSmall', parent=styles['Normal'], 
                                fontName='DejaVu', fontSize=9, leading=12, textColor=HexColor('#666666'))
    
    tracks = report_data.get('tracks', {})
    cons = report_data.get('consolidation', {})
    
    # Заголовок
    story.append(Paragraph('BizEval', title_style))
    story.append(Paragraph('Стратегический Анализ Бизнеса', h1_style))
    story.append(Spacer(1, 20))
    
    # Исходные данные
    if form_input:
        story.append(Paragraph('ИСХОДНЫЕ ДАННЫЕ ДЛЯ АНАЛИЗА', h2_style))
        story.append(Spacer(1, 10))
        
        fields = [
            ('Отрасль и продукты', 'industry_products'),
            ('Клиенты', 'customers'),
            ('Бизнес-модель', 'business_model'),
            ('География', 'geography'),
            ('Ограничения', 'constraints'),
            ('Стратегические цели', 'strategic_goals'),
            ('Дополнительно', 'additional_info')
        ]
        
        for label, key in fields:
            value = form_input.get(key, '')
            if value:
                story.append(Paragraph(f'<b>{label}:</b> {value}', small_style))
                story.append(Spacer(1, 5))
        
        story.append(Spacer(1, 20))
    
    # Трек 1
    story.append(Paragraph('АНАЛИЗ РЫНКОВ И НИШ', h1_style))
    story.append(Spacer(1, 10))
    market_text = tracks.get('market_analysis', 'Нет данных')
    for para in market_text.split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', ' '), text_style))
            story.append(Spacer(1, 8))
    story.append(Spacer(1, 12))
    
    # Трек 2
    story.append(Paragraph('АНАЛИЗ АНАЛОГОВ И АНТИЛОГОВ', h1_style))
    story.append(Spacer(1, 10))
    growth_text = tracks.get('growth_opportunities', 'Нет данных')
    for para in growth_text.split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', ' '), text_style))
            story.append(Spacer(1, 8))
    story.append(Spacer(1, 12))
    
    # Трек 3
    story.append(Paragraph('АНАЛИЗ КЛИЕНТСКИХ БОЛЕЙ', h1_style))
    story.append(Spacer(1, 10))
    risks_text = tracks.get('risks_constraints', 'Нет данных')
    for para in risks_text.split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', ' '), text_style))
            story.append(Spacer(1, 8))
    story.append(Spacer(1, 12))
    
    # Итоговое резюме
    story.append(PageBreak())
    story.append(Paragraph('ИТОГОВОЕ РЕЗЮМЕ', h1_style))
    story.append(Spacer(1, 10))
    exec_summary = cons.get('executive_summary', 'Нет данных')
    for para in exec_summary.split('\n\n'):
        if para.strip():
            story.append(Paragraph(para.strip().replace('\n', ' '), text_style))
            story.append(Spacer(1, 8))
    
    doc.build(story)
    return output_path


def generate_docx(report_data: dict, output_path: str, form_input: dict = None):
    doc = Document()
    tracks = report_data.get('tracks', {})
    cons = report_data.get('consolidation', {})
    
    # Заголовок
    title = doc.add_heading('BizEval', 0)
    title.runs[0].font.color.rgb = RGBColor(102, 126, 234)
    doc.add_heading('Стратегический Анализ Бизнеса', level=2)
    doc.add_paragraph('')
    
    # Исходные данные
    if form_input:
        doc.add_heading('ИСХОДНЫЕ ДАННЫЕ ДЛЯ АНАЛИЗА', level=2)
        
        fields = [
            ('Отрасль и продукты', 'industry_products'),
            ('Клиенты', 'customers'),
            ('Бизнес-модель', 'business_model'),
            ('География', 'geography'),
            ('Ограничения', 'constraints'),
            ('Стратегические цели', 'strategic_goals'),
            ('Дополнительно', 'additional_info')
        ]
        
        for label, key in fields:
            value = form_input.get(key, '')
            if value:
                p = doc.add_paragraph()
                run = p.add_run(f'{label}: ')
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)
                run = p.add_run(value)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(102, 102, 102)
        
        doc.add_paragraph('')
    
    # Трек 1
    doc.add_heading('АНАЛИЗ РЫНКОВ И НИШ', level=1)
    market_text = tracks.get('market_analysis', 'Нет данных')
    for para in market_text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph('')
    
    # Трек 2
    doc.add_heading('АНАЛИЗ АНАЛОГОВ И АНТИЛОГОВ', level=1)
    growth_text = tracks.get('growth_opportunities', 'Нет данных')
    for para in growth_text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph('')
    
    # Трек 3
    doc.add_heading('АНАЛИЗ КЛИЕНТСКИХ БОЛЕЙ', level=1)
    risks_text = tracks.get('risks_constraints', 'Нет данных')
    for para in risks_text.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    doc.add_paragraph('')
    
    # Итоговое резюме
    doc.add_page_break()
    doc.add_heading('ИТОГОВОЕ РЕЗЮМЕ', level=1)
    exec_summary = cons.get('executive_summary', 'Нет данных')
    for para in exec_summary.split('\n\n'):
        if para.strip():
            doc.add_paragraph(para.strip())
    
    doc.save(output_path)
    return output_path